"""Insert logbook workplace screenshots into existing report without changing other content."""
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement

SRC = Path(r"C:\Users\Leonm\Downloads\Laura_Abbey_Industrial_Attachment_Final_Report_WITH_SCREENSHOTS.docx")
OUT = Path(r"C:\Users\Leonm\Downloads\Laura_Abbey_Industrial_Attachment_Final_Report_WITH_SCREENSHOTS.docx")
LOGBOOK = Path(r"C:\Users\Leonm\Music\MarketSphere-repo\docs\logbook-images")

# Insert these logbook images at the start of each Appendix B week block
# (before the existing VS Code figures). Do not alter other document text.
INSERTS = {
    "Week 1: Foundation Networking and SAP MM Orientation": [
        (LOGBOOK / "week1/logbook_w1_1.jpeg", "Figure L1.1: Preparing Ethernet cable wires during RJ45 termination practice at BPC."),
        (LOGBOOK / "week1/logbook_w1_2.jpeg", "Figure L1.2: Networking workstation with T568A/T568B chart, RJ45 cable tester, crimping tool and punch down tools."),
    ],
    "Week 2: Automation Strategy, Migration Planning and DNS Checks": [
        (LOGBOOK / "week2/logbook_w2_1.jpeg", "Figure L2.1: SAP Create Purchase Requisition screen used during Week 2."),
        (LOGBOOK / "week2/logbook_w2_2.jpeg", "Figure L2.2: Power Automate cloud flow Summarize Uploaded Documents ready for testing."),
        (LOGBOOK / "week2/logbook_w2_3.jpeg", "Figure L2.3: SharePoint migration in progress for DocuSign Envelopes into the docusignenvelopes library."),
    ],
    "Week 3: Power BI Survey Dashboard and Power Automate Trigger Testing": [
        (LOGBOOK / "week3/logbook_w3_1.png", "Figure L3.1: BPC SHER Survey Responses Power BI dashboard."),
        (LOGBOOK / "week3/logbook_w3_2.png", "Figure L3.2: Power Automate cloud flow with SharePoint file created trigger, extract folder and apply to each steps."),
    ],
    "Week 4: SAP Integration Research and Safety Culture Dashboard QA": [
        (LOGBOOK / "week4/logbook_w4_1.png", "Figure L4.1: Near Miss Reporting Executive Safety Culture Overview Power BI dashboard."),
    ],
    "Week 5: KYC Support and SharePoint Desktop Flow Revamp": [
        (LOGBOOK / "week5/logbook_w5_1.png", "Figure L5.1: Power Automate Desktop flow Extract Summary File for DocuSign envelope summary processing."),
    ],
    "Week 6: DNS Server Upgrade from Windows Server 2016 to 2022": [
        (LOGBOOK / "week6/logbook_w6_1.jpeg", "Figure L6.1: Terminal on gab-cp-dns-02 showing Unbound DNS forwarder configuration edit and successful service start."),
    ],
    "Week 7: Veeam Restore Testing and 4G CPE Telephony Check": [
        (LOGBOOK / "week7/logbook_w7_1.png", "Figure L7.1: Veeam Restoring VM log for gab-cp-dns-02 failing at VM Encryption Policy after data restore."),
        (LOGBOOK / "week7/logbook_w7_2.png", "Figure L7.2: Veeam Backup and Replication console showing daily backup success and related job or SMTP warnings."),
        (LOGBOOK / "week7/logbook_w7_3.jpg", "Figure L7.3: 4G CPE admin page for the telephony support link showing RSRP, SINR and Band readings."),
    ],
}


def set_run_font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def make_caption_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True)
    return p


def make_image_paragraph(doc, image_path: Path, width_inches: float = 5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return p


def insert_after(anchor_para, new_para):
    """Move new_para XML to sit immediately after anchor_para."""
    anchor_para._element.addnext(new_para._element)


def main():
    doc = Document(str(SRC))

    # Build map of heading text -> paragraph
    headings = {}
    for para in doc.paragraphs:
        t = para.text.strip()
        if t in INSERTS:
            headings[t] = para

    missing = [k for k in INSERTS if k not in headings]
    if missing:
        # try partial match for week titles that may differ slightly
        for key in list(missing):
            for para in doc.paragraphs:
                if para.text.strip().startswith(key.split(":")[0]) and "Appendix" not in para.text:
                    # only appendix week headings are short titles without Week Ending
                    txt = para.text.strip()
                    if "Week Ending" in txt:
                        continue
                    if txt.startswith(key.split(":")[0] + ":"):
                        headings[key] = para
                        break
        missing = [k for k in INSERTS if k not in headings]

    if missing:
        print("Could not find headings:")
        for m in missing:
            print(" ", m)
        print("Available Week headings in appendix area:")
        for para in doc.paragraphs:
            if para.text.strip().startswith("Week ") and "Week Ending" not in para.text.strip():
                print(" ", repr(para.text.strip()))
        raise SystemExit(1)

    # Insert in reverse week order so earlier inserts don't shift later anchors badly
    # Actually we insert after heading, so order within week matters; do weeks from end to start
    for heading_text in reversed(list(INSERTS.keys())):
        anchor = headings[heading_text]
        items = INSERTS[heading_text]
        # Build image+caption pairs at end of doc first, then move after anchor in reverse
        # so final order is image1, caption1, image2, caption2...
        created = []
        for path, caption in items:
            if not path.exists():
                print("Missing image", path)
                continue
            img_p = make_image_paragraph(doc, path)
            cap_p = make_caption_paragraph(doc, caption)
            created.append((img_p, cap_p))

        # Move after anchor: insert last pair first so first pair ends up right after heading
        for img_p, cap_p in reversed(created):
            insert_after(anchor, cap_p)
            insert_after(anchor, img_p)

    # Remove empty leftover paragraphs at document end that python-docx leaves when we move nodes?
    # Not critical.

    try:
        doc.save(str(OUT))
        print(f"Saved {OUT}")
    except PermissionError:
        alt = OUT.with_name(OUT.stem + "_LOGBOOK_ADDED.docx")
        doc.save(str(alt))
        print(f"Target locked. Saved as {alt}")
        print("Close the original Word file and rename this over it if needed.")

    doc2 = Document(str(OUT if OUT.exists() else OUT.with_name(OUT.stem + "_LOGBOOK_ADDED.docx")))
    # recount - if save went to alt use that
    # Actually if PermissionError, OUT might be old. Check alt.
    saved = OUT if OUT.exists() else None
    # Verify by opening what we just wrote - if PermissionError path used alt
    print("Done.")


if __name__ == "__main__":
    main()
