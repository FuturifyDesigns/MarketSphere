"""Remove logbook Appendix A; rename screenshots appendix to Appendix A."""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DOC = r"C:\Users\Leonm\Downloads\Laura_Abbey_Industrial_Attachment_Final_Report_WITH_SCREENSHOTS_ALIGNED.docx"


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def replace_text_keep_style(para, new_text, *, size=None, bold=None, italic=None):
    # Prefer first run's formatting if present
    first = para.runs[0] if para.runs else None
    for r in para.runs:
        r.text = ""
    if first is not None:
        first.text = new_text
        if size is not None or bold is not None or italic is not None:
            set_run_font(
                first,
                size=size if size is not None else (first.font.size.pt if first.font.size else 12),
                bold=bold if bold is not None else bool(first.bold),
                italic=italic if italic is not None else bool(first.italic),
            )
    else:
        run = para.add_run(new_text)
        set_run_font(run, size=size or 12, bold=bool(bold), italic=bool(italic))


def main():
    doc = Document(DOC)

    # 1) Remove Appendix A logbook heading + placeholder (+ blank paras until Appendix B)
    to_remove = []
    removing = False
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("Appendix A: Weekly Logbooks"):
            removing = True
            to_remove.append(para)
            continue
        if removing:
            if t.startswith("Appendix B:"):
                removing = False
                break
            # remove placeholder / blanks between A and B
            to_remove.append(para)

    for para in to_remove:
        el = para._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # 2) Rename Appendix B -> Appendix A everywhere in paragraphs
    for para in doc.paragraphs:
        t = para.text
        if not t:
            continue
        new = t
        new = new.replace("Appendix B: Screenshots and Practical Evidence by Week",
                          "Appendix A: Screenshots and Practical Evidence by Week")
        new = new.replace("Appendix B contains", "Appendix A contains")
        new = new.replace("Appendix B under", "Appendix A under")
        new = new.replace("in Appendix B", "in Appendix A")
        new = new.replace("Appendix B", "Appendix A")  # catch remaining
        # Avoid double-fix if already A somehow - ok
        # Fix accidental "Appendix A: Weekly" shouldn't exist anymore
        if new != t:
            # Keep heading bold if it was a heading-like appendix title
            if new.startswith("Appendix A: Screenshots"):
                replace_text_keep_style(para, new, size=12, bold=True)
            else:
                replace_text_keep_style(para, new)

    # 3) Also fix section 12 intro if it still mentions logbooks only
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("12. APPENDICES"):
            continue
        if "signed weekly logbooks" in t.lower() and "Appendix" in t:
            replace_text_keep_style(
                para,
                "Appendix A contains the workplace screenshots and VS Code method notes for each week of the attachment.",
                size=12,
            )

    try:
        doc.save(DOC)
        print("Saved", DOC)
    except PermissionError:
        alt = DOC.replace(".docx", "_NO_LOGBOOK_APPENDIX.docx")
        doc.save(alt)
        print("Locked. Saved", alt)

    # verify
    doc2 = Document(DOC if True else DOC)
    try:
        doc2 = Document(DOC)
    except Exception:
        doc2 = Document(DOC.replace(".docx", "_NO_LOGBOOK_APPENDIX.docx"))

    print("--- Appendix headings ---")
    for p in doc2.paragraphs:
        t = p.text.strip()
        if t.startswith("Appendix"):
            print(t[:120])
        if "Appendix B" in t and not t.startswith("Appendix A"):
            print("LEFTOVER B:", t[:120])


if __name__ == "__main__":
    main()
