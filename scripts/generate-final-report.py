"""Generate CSI352/ISS302 Industrial Attachment Final Report
in the same format style as the Leon Maunge example report.
"""
from pathlib import Path

from report_weeks import add_weeks

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, Cm, RGBColor, Twips

OUT = Path(r"C:\Users\Leonm\Downloads\Laura_Abbey_Industrial_Attachment_Final_Report_WITH_SCREENSHOTS_UPDATED.docx")
LOGO = Path(r"C:\Users\Leonm\Music\MarketSphere-repo\docs\ub-logo.png")
SHOTS = Path(r"C:\Users\Leonm\Music\MarketSphere-repo\docs\report-screenshots")
LOGBOOK = Path(r"C:\Users\Leonm\Music\MarketSphere-repo\docs\logbook-images")

STUDENT_NAME = "LAURA ADEI RESEGO ABBEY"
STUDENT_ID = "202200191"
COURSE_CODE = "ISS302 & CSI352"
PROGRAMME = "BSc Computing with Finance"
HOST = "BOTSWANA POWER CORPORATION (BPC)"
HOST_DEPT = "[Department Attached to]"
IND_SUPERVISOR = "Mr Tefo Stanley Ramonoko"
IND_POSITION = "[Position]"
UNI_TUTOR = "[University Tutor Name]"
ATTACH_DATES = "2nd June to 16th July 2026"

# Appendix B: (week title, list of (path, caption))
# Workplace photos from weekly logbooks come first, then method screenshots.
APPENDIX_B = [
    (
        "Week 1: Foundation Networking and SAP MM Orientation",
        [
            (LOGBOOK / "week1/logbook_w1_1.jpeg", "Figure B1.1: Preparing Ethernet cable wires during RJ45 termination practice at BPC."),
            (LOGBOOK / "week1/logbook_w1_2.jpeg", "Figure B1.2: Networking workstation showing T568A/T568B wiring chart, RJ45 cable tester, crimping tool and punch down tools used in Week 1."),
            (SHOTS / "week1_ethernet_t568b.png", "Figure B1.3: VS Code helper used to record T568B pin order and cable tester pass or fail logic."),
            (SHOTS / "week1_sap_mm.png", "Figure B1.4: SAP MM orientation notes from Week 1."),
        ],
    ),
    (
        "Week 2: Automation Strategy, Migration Planning and SAP Purchase Requisitions",
        [
            (LOGBOOK / "week2/logbook_w2_1.jpeg", "Figure B2.1: SAP Create Purchase Requisition screen used while learning the ME51N style procurement path."),
            (LOGBOOK / "week2/logbook_w2_2.jpeg", "Figure B2.2: Power Automate cloud flow Summarize Uploaded Documents showing the flow ready for testing."),
            (LOGBOOK / "week2/logbook_w2_3.jpeg", "Figure B2.3: SharePoint migration in progress for DocuSign Envelopes from \\\\gab-cp-dcsign-1 to the SharePoint docusignenvelopes library."),
            (SHOTS / "week2_automation_decision.png", "Figure B2.4: Automation decision helper for the under five minutes rule and SPMT versus Power Automate choice."),
            (SHOTS / "week2_power_query.png", "Figure B2.5: Power Query M language notes used for cleaning survey data before Power BI."),
            (SHOTS / "week2_dns_connectivity.png", "Figure B2.6: PowerShell DNS and connectivity checks used during server and remote desktop troubleshooting."),
        ],
    ),
    (
        "Week 3: Power BI Survey Dashboard and Power Automate Trigger Testing",
        [
            (LOGBOOK / "week3/logbook_w3_1.png", "Figure B3.1: BPC SHER Survey Responses Power BI dashboard showing total responses, departments and question level visuals."),
            (LOGBOOK / "week3/logbook_w3_2.png", "Figure B3.2: Power Automate cloud flow design with SharePoint file created trigger, extract folder, get files and apply to each conditions."),
            (SHOTS / "week3_powerbi_survey.png", "Figure B3.3: VS Code notes on survey question grouping and sample DAX measures for the single page dashboard approach."),
            (SHOTS / "week3_power_automate_trigger.png", "Figure B3.4: Trigger troubleshooting checklist used after the cloud flow failed to fire as expected during testing."),
        ],
    ),
    (
        "Week 4: SAP Integration Research and Safety Culture Dashboard Finalisation",
        [
            (LOGBOOK / "week4/logbook_w4_1.png", "Figure B4.1: Near Miss Reporting Executive Safety Culture Overview Power BI dashboard, including Goodhart Risk Index, department detail and sample size caution areas."),
            (SHOTS / "week4_sap_integration.png", "Figure B4.2: SAP integration path decision framework comparing cloud connector, RPA and native SAP tools."),
            (SHOTS / "week4_goodhart_qa.png", "Figure B4.3: Goodhart risk sample size caution helper used during dashboard quality review."),
        ],
    ),
    (
        "Week 5: KYC Support and SharePoint Desktop Flow Revamp",
        [
            (LOGBOOK / "week5/logbook_w5_1.png", "Figure B5.1: Power Automate Desktop flow Extract Summary File showing library path, zip path, unzip, Summary*.pdf retrieval and PDF text extraction steps."),
            (SHOTS / "week5_sharepoint_flow.png", "Figure B5.2: VS Code outline of the same desktop flow logic for DocuSign envelope summary updates to SharePoint."),
        ],
    ),
    (
        "Week 6: DNS Server Upgrade and DNS Service Validation",
        [
            (LOGBOOK / "week6/logbook_w6_1.jpeg", "Figure B6.1: Terminal work on gab-cp-dns-02 showing Unbound DNS forwarder configuration edit and successful service start after an initial failure."),
            (SHOTS / "week6_dns_validation.png", "Figure B6.2: PowerShell validation commands for DNS zones, forwarders, reverse lookup and Active Directory SRV records."),
            (SHOTS / "week6_dns_runbook.png", "Figure B6.3: DNS migration runbook outline covering assess, plan, backup, build, migrate, test, cutover and document."),
        ],
    ),
    (
        "Week 7: Veeam Restore Testing and 4G CPE Telephony Check",
        [
            (LOGBOOK / "week7/logbook_w7_1.png", "Figure B7.1: Veeam Restoring VM log for gab-cp-dns-02 showing successful data restore and registration, then failure at VM Encryption Policy."),
            (LOGBOOK / "week7/logbook_w7_2.png", "Figure B7.2: Veeam Backup and Replication console showing BPC DAILY BACKUPS success, BACKUP-SAP failure status and SMTP report delivery warning."),
            (LOGBOOK / "week7/logbook_w7_3.jpg", "Figure B7.3: 4G CPE admin home page for the telephony support link, including RSRP, SINR and Band 40 readings, with monitoring tools open in other browser tabs."),
            (SHOTS / "week7_veeam_restore.png", "Figure B7.4: VS Code restore test checklist documenting isolated restore settings and follow up items."),
            (SHOTS / "week7_cpe_signal.png", "Figure B7.5: VS Code helper used to interpret CPE RSRP quality bands during the telephony check."),
        ],
    ),
]


def set_run_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text, *, size=12, bold=False, italic=False, align="left",
             space_after=8, space_before=0, first_line=True, center=False):
    p = doc.add_paragraph()
    if center or align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line and align == "justify":
        pf.first_line_indent = Cm(1.25)

    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading_text(doc, text, level=1):
    """Match example: numbered ALL CAPS main headings."""
    sizes = {0: 14, 1: 12, 2: 12, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0 if level == 0 else 14)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True)
    return p


def add_bullet(doc, text, *, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=size)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size)
    return p


def add_image(doc, path: Path, caption: str, width_inches: float = 5.8):
    if not path.exists():
        add_para(doc, f"[Missing image: {path.name}]", italic=True, first_line=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_para(doc, caption, size=11, italic=True, center=True, space_after=12, first_line=False)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_page_number(section):
    """Footer style similar to example: N | Page"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_before = p.add_run()
    set_run_font(run_before, size=11)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run_before._r.append(fld_begin)
    run2 = p.add_run()
    set_run_font(run2, size=11)
    run2._r.append(instr)
    run3 = p.add_run()
    set_run_font(run3, size=11)
    run3._r.append(fld_sep)
    run4 = p.add_run("1")
    set_run_font(run4, size=11)
    run5 = p.add_run()
    set_run_font(run5, size=11)
    run5._r.append(fld_end)

    run_after = p.add_run(" | Page")
    set_run_font(run_after, size=11)


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        add_page_number(section)


def cover_page(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.5))

    add_para(doc, "UNIVERSITY OF BOTSWANA", size=14, bold=True, center=True,
             space_before=10, space_after=6, first_line=False)
    add_para(doc, "DEPARTMENT OF COMPUTER SCIENCE", size=12, bold=True, center=True,
             space_after=18, first_line=False)
    add_para(doc, "INDUSTRIAL ATTACHMENT FINAL REPORT", size=13, bold=True, center=True,
             space_after=18, first_line=False)

    # Example-style labelled fields, left aligned block centered via wide margins feel
    fields = [
        f"Course Code: {COURSE_CODE}",
        f"Student ID: {STUDENT_ID}",
        f"Student Name: {STUDENT_NAME}",
        f"Programme of Study: {PROGRAMME}",
        f"Host Organization: {HOST}",
        f"Department Attached to: {HOST_DEPT}",
        f"Industrial Supervisor: {IND_SUPERVISOR}",
        f"Position: {IND_POSITION}",
        f"University Tutor: {UNI_TUTOR}",
        f"Attachment Period: {ATTACH_DATES}",
    ]
    for line in fields:
        add_para(doc, line, size=12, center=True, space_after=6, first_line=False)


def acknowledgements(doc):
    page_break(doc)
    add_heading_text(doc, "ACKNOWLEDGEMENTS", level=0)
    add_para(
        doc,
        "I would like to express my sincere gratitude to the University of Botswana, "
        "Department of Computer Science, for providing me with the opportunity to "
        "undertake this industrial attachment. My heartfelt appreciation goes to my "
        "industrial supervisor at Botswana Power Corporation, Mr Tefo Stanley Ramonoko, "
        "for guidance, mentorship and patience throughout the attachment period.",
        align="justify",
    )
    add_para(
        doc,
        "I am particularly grateful to my university tutor for continuous support and "
        "valuable feedback during the weekly logbook reviews. Special thanks to all the "
        "staff members at Botswana Power Corporation who welcomed me and shared their "
        "expertise across networking, SAP, Microsoft 365 automation, Power BI, Windows "
        "Server DNS work, Veeam backup testing and telephony related checks, making this "
        "learning experience both enriching and memorable.",
        align="justify",
    )
    add_para(
        doc,
        "Finally, I acknowledge my fellow students and family members who provided "
        "encouragement and support throughout this journey.",
        align="justify",
    )


def toc_page(doc):
    page_break(doc)
    add_heading_text(doc, "TABLE OF CONTENTS", level=0)

    entries = [
        ("1. Introduction", "4"),
        ("2. Statement of Place of Attachment", "5"),
        ("3. Summary of Industrial Attachment Organization Background Information", "6"),
        ("4. Summary of Tasks/Activities Done During Attachment", "7"),
        ("5. Opportunities and Skills Gained During Attachment", "11"),
        ("6. Contribution Student Made to Department During Attachment", "13"),
        ("7. Strengths and Weaknesses Student Experienced During Attachment", "14"),
        ("8. Limitations & Challenges", "15"),
        ("9. General Observation about Attachment Exercise", "16"),
        ("10. Recommendations", "17"),
        ("11. Conclusions", "18"),
        ("12. Appendices", "19"),
    ]

    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(14.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(f"{title}\t{page}")
        set_run_font(run, size=12)


def body(doc):
    # 1
    page_break(doc)
    add_heading_text(doc, "1. INTRODUCTION", level=0)
    add_para(
        doc,
        "The industrial attachment program serves as a crucial bridge between theoretical "
        "knowledge acquired during university studies and practical application in real "
        "world working environments. This attachment at Botswana Power Corporation "
        "provided an invaluable opportunity to gain hands on experience in networking, "
        "enterprise systems, automation, management reporting and infrastructure support "
        "within a professional utility information technology environment.",
        align="justify",
    )
    add_para(
        doc,
        "During this period from 2nd June to 16th July 2026, I was exposed to practical "
        "Ethernet cable termination and testing, SAP Materials Management navigation and "
        "purchase requisition processes, Microsoft 365 migration and Power Automate "
        "planning, Power BI dashboard development and quality review, SharePoint desktop "
        "automation linked to DocuSign envelope summaries, a DNS server upgrade from "
        "Windows Server 2016 to Windows Server 2022, and backup disaster recovery testing "
        "using Veeam Backup and Replication, including a check of telephony related 4G "
        "CPE hardware. The weekly logbooks form the factual basis of this report. The "
        "sections that follow summarise the host organisation, the tasks completed, the "
        "skills gained, the challenges faced and the influence of the attachment on my "
        "career expectations as a BSc Computing with Finance student.",
        align="justify",
    )

    # 2
    page_break(doc)
    add_heading_text(doc, "2. STATEMENT OF PLACE OF ATTACHMENT", level=0)
    add_para(
        doc,
        "Botswana Power Corporation served as my host organization for this industrial "
        "attachment. BPC is the national electricity utility of Botswana and depends on "
        "information technology for procurement support, document management, identity "
        "and name resolution services, backup and recovery, management reporting and "
        "day to day staff productivity. My industrial supervisor was Mr Tefo Stanley "
        "Ramonoko. The attachment period ran from early June 2026 through 16th July 2026 "
        "across seven weekly logbook periods.",
        align="justify",
    )
    add_para(
        doc,
        "Day to day work took place within the Corporation’s information technology "
        "environment rather than in one narrow task only. Activities ranged from physical "
        "networking and SAP related processes through Microsoft 365 automation and Power "
        "BI reporting, to infrastructure work involving Windows Server DNS upgrades, "
        "Veeam restore testing and telephony hardware checks. Where internal systems, "
        "hostnames or compliance records are confidential, this report describes methods "
        "and learning outcomes without repeating restricted production detail beyond what "
        "was already recorded in the weekly logbooks for academic evaluation.",
        align="justify",
    )
    add_para(
        doc,
        "The placement provided a practical opportunity to see how computing skills "
        "support a large utility. It also showed how finance related organisational "
        "processes such as purchase requisitions and compliance verification sit beside "
        "technical platforms, which is especially relevant to Computing with Finance as "
        "a programme of study.",
        align="justify",
    )

    # 3
    page_break(doc)
    add_heading_text(doc, "3. SUMMARY OF INDUSTRIAL ATTACHMENT ORGANIZATION BACKGROUND INFORMATION", level=0)
    add_para(
        doc,
        "Botswana Power Corporation is responsible for generating, transmitting and "
        "distributing electricity in Botswana. As a large public utility it must keep "
        "services available, manage large volumes of organisational data, support staff "
        "across business units and protect systems that affect both internal operations "
        "and customer facing services. Because of that role, information technology is "
        "central to procurement, document control, identity services, network name "
        "resolution, backup and disaster recovery, management reporting and staff "
        "productivity tools.",
        align="justify",
    )
    add_para(
        doc,
        "From what I observed during the attachment, IT at BPC uses a mix of enterprise "
        "and infrastructure platforms. SAP is used for materials and procurement related "
        "processes, including Materials Management and purchase requisition paths often "
        "associated with transaction ME51N. Microsoft 365 tools including SharePoint, "
        "Power Automate and Power BI support collaboration, workflow automation and "
        "management reporting. Active Directory and DNS provide domain and name "
        "resolution services. Veeam Backup and Replication protects virtual machines "
        "covering areas such as DNS servers, domain controllers, Exchange, file servers, "
        "vending systems, websites and SAP related workloads. Physical networking and "
        "telephony support remain part of the same picture, from RJ45 cable work and "
        "remote desktop access into data centre systems through to checks on a 4G CPE "
        "router that supports telephone system connectivity.",
        align="justify",
    )
    add_para(
        doc,
        "The research and operational environment emphasised careful testing, "
        "documentation and change control, especially during DNS migration and backup "
        "restore validation. Compliance related work such as Know Your Customer form "
        "verification also showed that accurate data is treated as an operational "
        "requirement, not only a technical preference. Overall, IT at BPC links business "
        "process systems, compliance activities and the platforms that keep services "
        "findable, recoverable and usable.",
        align="justify",
    )

    # 4
    page_break(doc)
    add_heading_text(doc, "4. SUMMARY OF TASKS/ACTIVITIES DONE DURING ATTACHMENT", level=0)
    add_para(
        doc,
        "This section summarises the tasks and activities completed during each week of "
        "the attachment, following the weekly logbooks. Each week ends with the "
        "screenshots that should be placed in Appendix B as evidence.",
        align="justify",
        first_line=False,
    )
    add_weeks(doc, add_heading_text, add_para, add_bullet)

    # 5
    page_break(doc)
    add_heading_text(doc, "5. OPPORTUNITIES AND SKILLS GAINED DURING ATTACHMENT", level=0)

    add_heading_text(doc, "Relevant University Course Preparation", level=2)
    add_para(
        doc,
        "Several university courses provided essential preparation for the technical "
        "challenges encountered during this attachment. Networking modules established "
        "crucial foundations for Ethernet termination, cable testing, IP related "
        "troubleshooting and DNS name resolution that became central to Weeks 1, 2 and "
        "6. Database and information systems learning supported Power Query cleaning, "
        "Power BI modelling and careful reading of survey and safety metrics. "
        "Programming and problem solving courses helped when designing automation "
        "decision rules, reading Power Automate conditions and breaking infrastructure "
        "work into ordered tested steps. Systems analysis style thinking helped when "
        "comparing SAP integration options and when writing rollback minded plans for "
        "DNS cutover. As a Computing with Finance student, exposure to SAP purchase "
        "requisitions and KYC verification also connected technical platforms to "
        "organisational control and compliance processes.",
        align="justify",
    )

    add_heading_text(doc, "Technical Skills Development", level=2)
    add_para(
        doc,
        "This attachment provided extensive exposure to practical networking, SAP MM "
        "navigation and purchase requisitions, SharePoint Migration Tool planning, Power "
        "Automate cloud and desktop flows, Power BI dashboard development and quality "
        "assurance, Windows Server DNS migration on an Active Directory integrated "
        "estate, and Veeam entire virtual machine restore testing. I also gained "
        "familiarity with cellular signal metrics such as RSRP and SINR while checking "
        "a 4G CPE device supporting telephony.",
        align="justify",
    )

    add_heading_text(doc, "Enterprise Software and Automation Practices", level=2)
    add_para(
        doc,
        "I gained a clearer understanding of when to automate, how to separate bulk "
        "migration from post migration enrichment, and how to choose between SAP cloud "
        "connectors, robotic process automation and native SAP tools using a reusable "
        "decision question. Testing a flow whose trigger did not fire taught me that "
        "build completion is not the same as operational readiness. Revamping the "
        "desktop flow for DocuSign related SharePoint summaries strengthened skills in "
        "path configuration, zip handling, filtered retrieval and conditional logic.",
        align="justify",
    )

    add_heading_text(doc, "Infrastructure, Backup and Operational Discipline", level=2)
    add_para(
        doc,
        "The DNS upgrade and Veeam restore work offered deep insight into change "
        "control, coexistence of old and new services, rollback readiness, isolated "
        "test restores and careful log reading when a job fails late. Reviewing daily "
        "backup success alongside SMTP reporting failure and a problematic SAP agent "
        "backup job reinforced proactive console based monitoring rather than depending "
        "only on email alerts.",
        align="justify",
    )

    # 6
    page_break(doc)
    add_heading_text(doc, "6. CONTRIBUTION STUDENT MADE TO DEPARTMENT DURING ATTACHMENT", level=0)

    add_heading_text(doc, "Operational and Project Support", level=2)
    add_para(
        doc,
        "During the attachment I contributed by completing assigned operational and "
        "project tasks that supported both daily IT work and longer term improvements. "
        "Early contributions included practical networking support through cable "
        "termination and testing, and building familiarity with SAP Materials Management "
        "and purchase requisition processes used by the organisation.",
        align="justify",
    )

    add_heading_text(doc, "Automation, Reporting and Decision Support", level=2)
    add_para(
        doc,
        "I helped define when to automate, documented why the SharePoint Migration Tool "
        "should carry bulk moves, prepared survey data, delivered Power BI "
        "visualisations and left a reusable SAP integration decision framework. I also "
        "reviewed and corrected the Near Miss Reporting Executive Safety Culture "
        "Overview dashboard so leadership could use a cleaner view of reporting "
        "behaviour and Goodhart risk, with clear flags where high scores rested on low "
        "response counts.",
        align="justify",
    )

    add_heading_text(doc, "Compliance Support, Desktop Automation and Infrastructure", level=2)
    add_para(
        doc,
        "Later contributions included helping verify KYC form information, improving "
        "the desktop automation that updates SharePoint Content Summary values from "
        "summary PDFs inside DocuSign envelope packages, taking part in the DNS upgrade "
        "from Windows Server 2016 to Windows Server 2022, and carrying out a structured "
        "Veeam entire virtual machine restore test for a DNS server backup. The restore "
        "test did not fully succeed because of the VM Encryption Policy mismatch, but "
        "documenting the exact failure point, successful daily backup status, SMTP "
        "reporting faults and the borderline CPE signal reading still gave the "
        "department actionable follow up items.",
        align="justify",
    )

    # 7
    page_break(doc)
    add_heading_text(doc, "7. STRENGTHS AND WEAKNESSES STUDENT EXPERIENCED DURING ATTACHMENT", level=0)

    add_heading_text(doc, "Strengths Demonstrated", level=2)
    add_para(
        doc,
        "One strength I experienced was the ability to learn unfamiliar enterprise tools "
        "quickly when guided. SAP navigation, Power BI modelling, Power Automate flow "
        "testing and the Veeam restore wizard were all new in depth, yet I could move "
        "from observation to useful participation within the same week. Another strength "
        "was becoming more systematic about testing and documentation. When the Power "
        "Automate trigger failed, when dashboard metrics looked extreme, and when the "
        "Veeam restore stopped at encryption policy, I focused on identifying the exact "
        "point of failure instead of only saying that the task did not work. Linking "
        "university theory to workplace practice also came naturally, for example "
        "relating cabling standards to tester results, or relating name resolution theory "
        "to real remote desktop and DNS faults.",
        align="justify",
    )

    add_heading_text(doc, "Areas for Improvement", level=2)
    add_para(
        doc,
        "Weaknesses were also clear. I had limited prior exposure to some BPC specific "
        "platforms, so early progress on SAP, Power Platform and Veeam depended heavily "
        "on supervisor and staff guidance. At times I underestimated environment "
        "dependencies such as connector permissions, on premises gateway requirements "
        "for SAP integration, or datastore encryption policies during restore. I also "
        "saw that building a flow or dashboard is not enough if trigger paths, sample "
        "sizes and restore policies are not validated. To improve, I need to ask earlier "
        "about infrastructure prerequisites and to use written test checklists before "
        "calling a task complete.",
        align="justify",
    )

    # 8
    page_break(doc)
    add_heading_text(doc, "8. LIMITATIONS & CHALLENGES", level=0)

    add_heading_text(doc, "Confidentiality and Evidence Limits", level=2)
    add_para(
        doc,
        "Access and confidentiality limited what could be copied out of the "
        "organisation. The Week 6 logbook notes that some infrastructure screenshots "
        "could not be shared, so parts of the evidence pack must remain redacted or "
        "replaced with approved summaries. This limitation affected how completely "
        "live production evidence can appear in Appendix B, even though the methods and "
        "outcomes are described in the weekly narrative.",
        align="justify",
    )

    add_heading_text(doc, "Automation and Integration Challenges", level=2)
    add_para(
        doc,
        "The Power Automate cloud flow built earlier failed testing because the trigger "
        "did not fire, even though the flow design itself had been completed. "
        "Integrating cloud triggers with desktop file processing required careful "
        "library paths, zip extraction and a condition that stopped the flow when no "
        "Summary*.pdf file was present. Without those guards the SharePoint Content "
        "Summary update would have been unreliable.",
        align="justify",
    )

    add_heading_text(doc, "Infrastructure and Restore Challenges", level=2)
    add_para(
        doc,
        "DNS migration required old and new servers to coexist, forwarders to be copied "
        "accurately, reverse zones and SRV records to be validated, and a rollback path "
        "to remain ready until the new Windows Server 2022 service proved stable. During "
        "Veeam testing, an entire virtual machine restore transferred data and "
        "registered the machine but failed at VM Encryption Policy application. At the "
        "same time, backup email reporting was broken by invalid SMTP credentials, and "
        "a SAP related agent backup job was in a failed or disabled state for an "
        "extended period. These issues were approached by isolating test changes from "
        "production, reading logs to the failing step, documenting findings and "
        "escalating configuration items that needed administrator rights.",
        align="justify",
    )

    # 9
    page_break(doc)
    add_heading_text(doc, "9. GENERAL OBSERVATION ABOUT ATTACHMENT EXERCISE", level=0)

    add_heading_text(doc, "Utility IT as Both Technical and Procedural Work", level=2)
    add_para(
        doc,
        "My general observation is that the attachment exercise is one of the most "
        "effective parts of the computing programme when the host organisation gives "
        "students real tasks. At BPC the work was not limited to watching. I could "
        "terminate cables, navigate SAP, plan automation, build and correct dashboards, "
        "improve a desktop flow, take part in a DNS upgrade and run a structured restore "
        "test. That mix showed that IT in a utility is both technical and procedural. A "
        "correct command matters, but so do rollback plans, sample size caution on "
        "management reports, and isolation of test restores from live networks.",
        align="justify",
    )

    add_heading_text(doc, "Value of Weekly Reflection and Professional Expectations", level=2)
    add_para(
        doc,
        "I also observed that weekly reflection through logbooks is valuable. Writing "
        "each week forced me to separate what I did, what was interesting, what I "
        "learnt and what remained outstanding. Working under Mr Ramonoko and within BPC "
        "processes gave me a clearer picture of professional expectations around "
        "accuracy, compliance and change control than lectures alone had provided. The "
        "attachment also showed how computing work supports finance related controls "
        "such as purchase requisitions and KYC verification, which strengthens the "
        "relevance of Computing with Finance as a programme path.",
        align="justify",
    )

    # 10
    page_break(doc)
    add_heading_text(doc, "10. RECOMMENDATIONS", level=0)

    add_heading_text(doc, "For Future Students", level=2)
    add_para(
        doc,
        "Future students undertaking similar attachments should strengthen foundational "
        "networking and systems troubleshooting skills before the placement, as this "
        "allows more time to focus on enterprise tools and infrastructure projects. "
        "Students should also maintain detailed weekly logbooks and actively seek "
        "complete task lifecycles from planning through testing and documentation. I "
        "recommend clarifying early which screenshots may be restricted by "
        "confidentiality so that approved evidence for the final report can be planned "
        "in good time.",
        align="justify",
    )

    add_heading_text(doc, "For University Curriculum", level=2)
    add_para(
        doc,
        "The university curriculum could benefit from increased practical exposure to "
        "enterprise systems concepts, basic cabling and DNS labs, data preparation for "
        "reporting, and introductory workflow automation ideas. Short project based "
        "exercises that simulate change control, rollback planning and restore testing "
        "would better prepare students for utility and enterprise IT environments. "
        "Additional bridging between computing topics and organisational controls would "
        "also help Computing with Finance students see how technical platforms support "
        "procurement and compliance processes.",
        align="justify",
    )

    add_heading_text(doc, "For the Attachment Process and Host Department", level=2)
    add_para(
        doc,
        "A short induction checklist covering the main systems a student will touch, "
        "including access requests and data centre safety rules, would save time. Where "
        "possible, keeping one end to end project visible while still rotating through "
        "support tasks helps learning stay coherent. Issues discovered during student "
        "testing should be tracked to closure, including the Power Automate trigger "
        "fault, the VM Encryption Policy mismatch on restore, the SMTP credential "
        "problem affecting backup reports, the long failing SAP agent backup job, and "
        "monitoring of the borderline 4G CPE signal supporting telephony.",
        align="justify",
    )

    # 11
    page_break(doc)
    add_heading_text(doc, "11. CONCLUSIONS", level=0)
    add_para(
        doc,
        "This industrial attachment at Botswana Power Corporation provided an "
        "exceptional opportunity to bridge the gap between academic learning and "
        "professional information technology practice. The experience encompassed "
        "exposure to networking, enterprise applications, automation, management "
        "reporting, Windows Server DNS migration and backup restore validation that "
        "significantly enhanced my understanding of how computing supports a national "
        "utility.",
        align="justify",
    )
    add_para(
        doc,
        "The progression from Ethernet cabling and SAP orientation through Power "
        "Platform and Power BI work to infrastructure cutover and disaster recovery "
        "testing provided a comprehensive overview of modern IT support and improvement "
        "work. Working with tools such as SAP, SharePoint Migration Tool, Power "
        "Automate, Power BI, Windows Server DNS and Veeam offered valuable insight into "
        "platforms that organisations actually depend on. The attachment also confirmed "
        "the importance of careful testing, documentation and professional discretion "
        "when production systems and compliance data are involved.",
        align="justify",
    )
    add_para(
        doc,
        "Most significantly, the attachment provided clarity regarding career direction "
        "and professional development goals. I leave with stronger interest in roles that "
        "combine systems administration, enterprise applications and process improvement "
        "through automation and reporting, while respecting change control. As a BSc "
        "Computing with Finance student, seeing purchase requisitions, KYC verification "
        "and executive safety reporting beside technical platforms helped confirm that "
        "computing skills and organisational controls belong together. This attachment "
        "has been instrumental in preparing for the transition from academic study to "
        "professional practice.",
        align="justify",
    )

    # 12
    page_break(doc)
    add_heading_text(doc, "12. APPENDICES", level=0)
    add_heading_text(doc, "Appendix A: Weekly Logbooks", level=2)
    add_para(
        doc,
        "[Signed copies of all weekly logbooks from Week 1 through Week 7]",
        align="justify",
        first_line=False,
    )

    page_break(doc)
    add_heading_text(doc, "Appendix B: Screenshots and Practical Evidence by Week", level=2)
    add_para(
        doc,
        "The figures below combine workplace evidence taken from the weekly logbooks "
        "with VS Code method screenshots. Logbook photos and system screens come first "
        "in each week, followed by helper scripts used to document the same work. Where "
        "live production detail is sensitive, captions describe the method without "
        "repeating unnecessary personal or confidential values.",
        align="justify",
        first_line=False,
    )

    for week_title, figures in APPENDIX_B:
        add_heading_text(doc, week_title, level=2)
        for path, caption in figures:
            add_image(doc, path, caption)

    page_break(doc)
    add_heading_text(doc, "Appendix C: Other Project Graphics or Notes", level=2)
    add_para(
        doc,
        "Attach any additional approved diagrams, decision notes, dashboard exports or "
        "migration summaries produced during the attachment, subject to confidentiality "
        "rules. Do not include passwords, personal KYC data or unrestricted internal "
        "network documentation.",
        align="justify",
        first_line=False,
    )


def main():
    doc = Document()
    set_doc_defaults(doc)
    cover_page(doc)
    acknowledgements(doc)
    toc_page(doc)
    body(doc)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
