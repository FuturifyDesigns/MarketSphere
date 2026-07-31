"""Align Appendix B captions and add signature spaces. No full rewrite of report body."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DOC = Path(r"C:\Users\Leonm\Downloads\Laura_Abbey_Industrial_Attachment_Final_Report_WITH_SCREENSHOTS.docx")

CAPTION_MAP = {
    "Figure L1.1: Preparing Ethernet cable wires during RJ45 termination practice at BPC.":
        "Figure L1.1 (workplace): Preparing Ethernet cable wires during RJ45 termination practice at BPC. Matches the cable wiring method notes in Figure B1.1.",
    "Figure L1.2: Networking workstation with T568A/T568B chart, RJ45 cable tester, crimping tool and punch down tools.":
        "Figure L1.2 (workplace): Networking workstation with T568A/T568B chart, RJ45 cable tester, crimping tool and punch down tools used for the same termination and testing work as Figure B1.1.",
    "Figure B1.1: Ethernet T568B wiring check script used to document RJ45 termination and cable tester logic.":
        "Figure B1.1 (VS Code method notes): T568B pin order and cable tester pass or fail checklist documenting the same RJ45 work shown in Figures L1.1 and L1.2.",
    "Figure B1.2: SAP MM orientation notes showing module purpose and practices during Week 1.":
        "Figure B1.2 (VS Code method notes): SAP Materials Management orientation notes from the same Week 1 SAP introduction described in Section 4.",

    "Figure L2.1: SAP Create Purchase Requisition screen used during Week 2.":
        "Figure L2.1 (workplace): SAP Create Purchase Requisition screen used while practising the procurement path described in Week 2.",
    "Figure L2.2: Power Automate cloud flow Summarize Uploaded Documents ready for testing.":
        "Figure L2.2 (workplace): Power Automate cloud flow Summarize Uploaded Documents ready for testing. This is the post migration automation approach summarised in Figure B2.1.",
    "Figure L2.3: SharePoint migration in progress for DocuSign Envelopes into the docusignenvelopes library.":
        "Figure L2.3 (workplace): SharePoint bulk migration of DocuSign Envelopes into the docusignenvelopes library. This is the bulk move approach that Figure B2.1 says should use SPMT rather than Power Automate.",
    "Figure B2.1: Automation decision helper for the under five minutes rule and SPMT versus Power Automate choice.":
        "Figure B2.1 (VS Code method notes): Decision helper for when to automate and when to use SPMT for bulk migrate versus Power Automate after the move. Explains the approach behind Figures L2.2 and L2.3.",
    "Figure B2.2: Power Query M language notes used for cleaning and standardising survey data before Power BI.":
        "Figure B2.2 (VS Code method notes): Power Query cleaning steps used to prepare survey data before the Power BI dashboards built in Week 3 (see Figures L3.1 and B3.1).",
    "Figure B2.3: PowerShell DNS and connectivity checks used during server and remote desktop troubleshooting.":
        "Figure B2.3 (VS Code method notes): DNS and connectivity PowerShell checks used during Week 2 server troubleshooting, and later related to the DNS validation work in Week 6.",

    "Figure L3.1: BPC SHER Survey Responses Power BI dashboard.":
        "Figure L3.1 (workplace): BPC SHER Survey Responses Power BI dashboard built in Week 3. Matches the survey grouping and measure notes in Figure B3.1.",
    "Figure L3.2: Power Automate cloud flow with SharePoint file created trigger, extract folder and apply to each steps.":
        "Figure L3.2 (workplace): Power Automate cloud flow with SharePoint file created trigger, extract folder and apply to each steps. Figure B3.2 records the trigger troubleshooting done when testing this kind of flow.",
    "Figure B3.1: Power BI style DAX notes and survey question grouping for the single page dashboard approach.":
        "Figure B3.1 (VS Code method notes): Survey question grouping and sample measures used for the single page dashboard approach shown in Figure L3.1.",
    "Figure B3.2: Power Automate trigger troubleshooting checklist used after the cloud flow failed to fire during testing.":
        "Figure B3.2 (VS Code method notes): Trigger troubleshooting checklist used when the cloud flow was built successfully but did not fire as expected during testing of flows such as Figure L3.2.",

    "Figure L4.1: Near Miss Reporting Executive Safety Culture Overview Power BI dashboard.":
        "Figure L4.1 (workplace): Near Miss Reporting Executive Safety Culture Overview dashboard finalised in Week 4. Figure B4.2 records the Goodhart risk and sample size checks applied to this dashboard.",
    "Figure B4.1: SAP integration path decision framework comparing cloud connector, RPA and native SAP tools.":
        "Figure B4.1 (VS Code method notes): SAP integration decision framework from the Week 4 research comparing Power Automate cloud connector, desktop RPA and native SAP tools.",
    "Figure B4.2: Goodhart risk sample size caution helper used while reviewing the safety culture Power BI dashboard.":
        "Figure B4.2 (VS Code method notes): Goodhart risk and low sample size caution helper used while quality checking the dashboard in Figure L4.1.",

    "Figure L5.1: Power Automate Desktop flow Extract Summary File for DocuSign envelope summary processing.":
        "Figure L5.1 (workplace): Power Automate Desktop flow Extract Summary File used to unzip DocuSign envelopes, find Summary*.pdf files and extract text for SharePoint. Figure B5.1 documents the same logic.",
    "Figure B5.1: Desktop flow logic for unzipping DocuSign envelopes, finding Summary*.pdf files and updating SharePoint Content Summary.":
        "Figure B5.1 (VS Code method notes): Step by step logic of the desktop flow shown in Figure L5.1 for path setup, unzip, Summary*.pdf filter, text clean and SharePoint Content Summary update.",

    "Figure L6.1: Terminal on gab-cp-dns-02 showing Unbound DNS forwarder configuration edit and successful service start.":
        "Figure L6.1 (workplace): DNS service work on gab-cp-dns-02 showing Unbound forwarder configuration edit and successful service start after an initial failure. Part of the same Week 6 DNS upgrade and validation effort as Figures B6.1 and B6.2.",
    "Figure B6.1: PowerShell validation commands for DNS zones, forwarders, reverse lookup and Active Directory SRV records.":
        "Figure B6.1 (VS Code method notes): Validation commands for DNS zones, forwarders, reverse lookup and Active Directory SRV records used to check the DNS environment during the Week 6 upgrade work related to Figure L6.1.",
    "Figure B6.2: DNS migration runbook outline covering assess, plan, backup, build, migrate, test, cutover and document.":
        "Figure B6.2 (VS Code method notes): DNS migration runbook outline for the Week 6 upgrade stages. Provides the structured plan behind the live DNS service work in Figure L6.1.",

    "Figure L7.1: Veeam Restoring VM log for gab-cp-dns-02 failing at VM Encryption Policy after data restore.":
        "Figure L7.1 (workplace): Veeam restore log for gab-cp-dns-02 showing successful data transfer and registration, then failure at VM Encryption Policy. Figure B7.1 records the same restore test plan and outcome.",
    "Figure L7.2: Veeam Backup and Replication console showing daily backup success and related job or SMTP warnings.":
        "Figure L7.2 (workplace): Veeam Backup and Replication console showing BPC DAILY BACKUPS success and related backup or SMTP reporting issues noted during the same Week 7 review as Figure L7.1.",
    "Figure L7.3: 4G CPE admin page for the telephony support link showing RSRP, SINR and Band readings.":
        "Figure L7.3 (workplace): 4G CPE admin page for the telephony support link showing RSRP, SINR and Band readings. Figure B7.2 interprets the same signal quality check.",
    "Figure B7.1: Veeam entire VM restore test checklist for gab-cp-dns-02-Test with network disconnected and encryption policy follow up.":
        "Figure B7.1 (VS Code method notes): Restore test checklist for gab-cp-dns-02-Test with network disconnected, matching the failed encryption policy restore shown in Figure L7.1 and the console review in Figure L7.2.",
    "Figure B7.2: 4G CPE signal quality helper used when reviewing the telephony backup link RSRP and SINR readings.":
        "Figure B7.2 (VS Code method notes): Signal quality helper used to interpret the RSRP and SINR readings from the CPE page in Figure L7.3.",
}

# Also match truncated/partial captions by prefix
CAPTION_PREFIXES = [
    ("Figure L1.1:", CAPTION_MAP[list(CAPTION_MAP.keys())[0]]),
]

WEEK_BLURBS = {
    "Week 1: Foundation Networking and SAP MM Orientation":
        "Week 1 evidence: workplace Ethernet termination photos first, then VS Code method notes for the same cabling checks, plus SAP MM orientation notes from the same week.",
    "Week 2: Automation Strategy, Migration Planning and DNS Checks":
        "Week 2 evidence: workplace SAP purchase requisition, Power Automate summarise flow and DocuSign SharePoint migration first, then VS Code notes for the automation decision rule, Power Query prep and DNS checks from the same planning week.",
    "Week 3: Power BI Survey Dashboard and Power Automate Trigger Testing":
        "Week 3 evidence: workplace Power BI survey dashboard and Power Automate flow first, then VS Code notes that document the same dashboard modelling and trigger troubleshooting.",
    "Week 4: SAP Integration Research and Safety Culture Dashboard QA":
        "Week 4 evidence: workplace executive safety culture dashboard first, then VS Code notes for the SAP integration research and the Goodhart sample size checks applied to that dashboard.",
    "Week 5: KYC Support and SharePoint Desktop Flow Revamp":
        "Week 5 evidence: workplace Power Automate Desktop Extract Summary File flow first, then VS Code notes describing the same unzip, Summary*.pdf and SharePoint update logic.",
    "Week 6: DNS Server Upgrade from Windows Server 2016 to 2022":
        "Week 6 evidence: workplace DNS service troubleshooting on gab-cp-dns-02 first, then VS Code validation commands and migration runbook notes for the same DNS upgrade week.",
    "Week 7: Veeam Restore Testing and 4G CPE Telephony Check":
        "Week 7 evidence: workplace Veeam restore failure log, backup console and 4G CPE page first, then VS Code checklists that record the same restore isolation settings and CPE signal interpretation.",
}

APPENDIX_INTRO_NEW = (
    "Appendix B contains two matching kinds of evidence for each week. Figures labelled L "
    "are workplace photos and system screens taken from the weekly logbooks. Figures "
    "labelled B are VS Code method notes and helper scripts for the same week’s tasks. "
    "Read each week as one set: the workplace screen shows what was done, and the VS Code "
    "figure explains the related method or checklist."
)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def replace_para_text(para, new_text, *, size=12, bold=False, italic=False, center=False):
    for r in para.runs:
        r.text = ""
    if para.runs:
        para.runs[0].text = new_text
        run = para.runs[0]
    else:
        run = para.add_run(new_text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def new_paragraph(doc, text, *, size=12, bold=False, italic=False, center=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def resolve_caption(text: str):
    if text in CAPTION_MAP:
        return CAPTION_MAP[text]
    # fuzzy: match by figure id prefix e.g. "Figure L1.1:"
    for old, new in CAPTION_MAP.items():
        old_id = old.split(":", 1)[0]
        if text.startswith(old_id + ":"):
            return new
    return None


def main():
    doc = Document(str(DOC))

    # Snapshot paragraph elements once
    paras = list(doc.paragraphs)

    # 1) Appendix intro
    for para in paras:
        t = para.text.strip()
        if t.startswith("The figures below") or t.startswith("Appendix B contains"):
            replace_para_text(para, APPENDIX_INTRO_NEW, size=12)
            break

    # 2) Captions
    updated = 0
    for para in paras:
        t = para.text.strip()
        if not t.startswith("Figure "):
            continue
        new = resolve_caption(t)
        if new:
            replace_para_text(para, new, size=11, italic=True, center=True)
            updated += 1

    # 3) Week blurbs using element addnext, no .index()
    for para in reversed(paras):
        t = para.text.strip()
        if t not in WEEK_BLURBS:
            continue
        nxt = para._element.getnext()
        nxt_text = ""
        if nxt is not None and nxt.tag.endswith("}p"):
            # get text from next paragraph element
            nxt_text = "".join(nxt.itertext()).strip()
        blurb = WEEK_BLURBS[t]
        if nxt_text.startswith("Week ") and "evidence:" in nxt_text.lower():
            # unusual
            pass
        if "evidence:" in nxt_text.lower():
            # update existing blurb paragraph
            for p in doc.paragraphs:
                if p._element is nxt:
                    replace_para_text(p, blurb, size=11, italic=True)
                    break
        else:
            new_p = new_paragraph(doc, blurb, size=11, italic=True, space_after=8)
            para._element.addnext(new_p._element)

    # 4) Signatures after Attachment Period
    for para in list(doc.paragraphs):
        if para.text.strip().startswith("Attachment Period:"):
            # check following few for existing signature
            el = para._element
            already = False
            cur = el.getnext()
            for _ in range(10):
                if cur is None:
                    break
                txt = "".join(cur.itertext()) if cur.tag.endswith("}p") else ""
                if "Student Signature" in txt:
                    already = True
                    break
                cur = cur.getnext()
            if not already:
                lines = [
                    " ",
                    "Student Signature: _______________________________     Date: ______________",
                    "Industrial Supervisor Signature: __________________     Date: ______________",
                    "University Tutor Signature: _______________________     Date: ______________",
                ]
                for line in reversed(lines):
                    new_p = new_paragraph(doc, line if line.strip() else " ", size=12, center=True, space_after=6)
                    para._element.addnext(new_p._element)
            break

    out = DOC
    try:
        doc.save(str(out))
        print(f"Saved {out}")
    except PermissionError:
        out = DOC.with_name(DOC.stem + "_ALIGNED.docx")
        doc.save(str(out))
        print(f"Locked. Saved {out}")

    # verify
    doc2 = Document(str(out))
    print("captions updated:", updated)
    print("has student signature:", any("Student Signature" in p.text for p in doc2.paragraphs))
    print("has workplace label:", any("(workplace)" in p.text for p in doc2.paragraphs))
    print("has VS Code label:", any("(VS Code method notes)" in p.text for p in doc2.paragraphs))


if __name__ == "__main__":
    main()
