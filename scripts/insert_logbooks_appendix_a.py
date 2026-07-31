"""Insert weekly logbook DOCX content into Appendix A of the aligned report."""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import tempfile
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

REPORT = Path(
    r"C:\Users\Leonm\Downloads\Laura_Abbey_Industrial_Attachment_Final_Report_WITH_SCREENSHOTS_ALIGNED.docx"
)
OUT = REPORT  # edit this document

LOGBOOKS = [
    (1, Path(r"d:\LogBook Week 1.docx")),
    (2, Path(r"d:\LogBook Week 2.docx")),
    (3, Path(r"d:\LogBook Week 3.docx")),
    (4, Path(r"d:\LOGBOOK WEEK 4..docx")),
    (5, Path(r"d:\LOGBOOK WEEK 5..docx")),
    (6, Path(r"d:\LOGBOOK WEEK 6..docx")),
    (7, Path(r"d:\LOGBOOK WEEK 7..docx")),
]


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def make_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)
    return p


def make_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True)
    return p


def page_break_para(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def copy_image_relationships(src_doc, dst_doc, element):
    """Ensure images referenced in element exist in dst_doc; rewrite rIds."""
    from io import BytesIO

    blips = element.xpath(".//*[local-name()='blip']")
    for blip in blips:
        embed = blip.get(qn("r:embed"))
        if not embed:
            continue
        try:
            rel = src_doc.part.rels[embed]
        except KeyError:
            continue
        blob = rel.target_part.blob
        new_rId, _img_part = dst_doc.part.get_or_add_image(BytesIO(blob))
        blip.set(qn("r:embed"), new_rId)

def append_logbook_body(dst_doc, src_path: Path, week: int, insert_before_el):
    """Copy body children from src logbook into dst, before insert_before_el."""
    src = Document(str(src_path))

    # Week separator heading first
    heading = make_heading(dst_doc, f"Logbook Week {week}")
    insert_before_el.addprevious(heading._element)

    note = make_note(
        dst_doc,
        f"Signed weekly logbook entry for Week {week}, as completed during the industrial attachment.",
    )
    insert_before_el.addprevious(note._element)

    body = src.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag in {"sectPr"}:
            continue
        new_el = deepcopy(child)
        # Remap images
        try:
            copy_image_relationships(src, dst_doc, new_el)
        except Exception as e:
            print(f"  week {week} image remap warning: {e}")
        insert_before_el.addprevious(new_el)

    # page break after each logbook
    pb = page_break_para(dst_doc)
    insert_before_el.addprevious(pb._element)
    print(f"Inserted Week {week} from {src_path.name}")


def main():
    doc = Document(str(REPORT))

    # Find Appendix A placeholder and Appendix B heading
    placeholder = None
    appendix_b = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("[Signed copies of all weekly logbooks"):
            placeholder = p
        if t.startswith("Appendix B:"):
            appendix_b = p
            break

    if placeholder is None:
        raise SystemExit("Appendix A placeholder not found")
    if appendix_b is None:
        raise SystemExit("Appendix B heading not found")

    # Replace placeholder text
    for r in placeholder.runs:
        r.text = ""
    if placeholder.runs:
        placeholder.runs[0].text = (
            "The signed weekly logbooks for Week 1 through Week 7 are included below."
        )
        set_run_font(placeholder.runs[0], size=12)
    else:
        run = placeholder.add_run(
            "The signed weekly logbooks for Week 1 through Week 7 are included below."
        )
        set_run_font(run, size=12)

    # Insert logbooks before Appendix B
    target = appendix_b._element
    for week, path in LOGBOOKS:
        if not path.exists():
            print("MISSING", path)
            continue
        append_logbook_body(doc, path, week, target)

    try:
        doc.save(str(OUT))
        print(f"Saved {OUT}")
    except PermissionError:
        alt = OUT.with_name(OUT.stem + "_WITH_LOGBOOKS.docx")
        doc.save(str(alt))
        print(f"Locked. Saved {alt}")


if __name__ == "__main__":
    main()
